"""
Document ingestion service for parsing and chunking
"""
import logging
from typing import List, Dict, Any
from uuid import UUID, uuid4
import io
import tempfile
import os

from pypdf import PdfReader
from docx import Document
import openpyxl
import tiktoken

from app.config import settings

from docling.document_converter import DocumentConverter, PdfFormatOption, WordFormatOption, CsvFormatOption, ExcelFormatOption, PowerpointFormatOption, MarkdownFormatOption, HTMLFormatOption, ImageFormatOption
from docling.datamodel.base_models import InputFormat
from docling.backend.pypdfium2_backend import PyPdfiumDocumentBackend
from docling.backend.msword_backend import MsWordDocumentBackend
from docling.backend.mspowerpoint_backend import MsPowerpointDocumentBackend
from docling.backend.msexcel_backend import MsExcelDocumentBackend
from docling.backend.csv_backend import CsvDocumentBackend
from docling.backend.image_backend import ImageDocumentBackend
from docling.backend.html_backend import HTMLDocumentBackend
from docling.backend.md_backend import MarkdownDocumentBackend
from docling.pipeline.simple_pipeline import SimplePipeline
from docling.pipeline.standard_pdf_pipeline import StandardPdfPipeline
from docling.datamodel.pipeline_options import PdfPipelineOptions

logger = logging.getLogger(__name__)


class IngestionService:
    """Service for document parsing and intelligent chunking"""
    
    def __init__(self):
        # Initialize tiktoken for accurate token counting
        self.tokenizer = tiktoken.get_encoding("cl100k_base")
    
    async def process_document(
        self,
        content: bytes,
        filename: str,
        file_type: str,
        resource_id: UUID,
        workspace_id: UUID
    ) -> List[Dict[str, Any]]:
        """
        Process a document into chunks
        
        Args:
            content: Raw file content
            filename: Original filename
            file_type: File extension
            resource_id: Resource UUID
            workspace_id: Workspace UUID
            
        Returns:
            List of chunk dictionaries
        """
        logger.info(f"Processing document: {filename} ({file_type})")
        
        # Extract text based on file type
        if file_type == "pdf":
            text = self._extract_pdf(content)
        elif file_type in ["docx", "doc"]:
            text = self._extract_docx(content)
        elif file_type in ["xlsx", "xls"]:
            text = self._extract_xlsx(content)
        elif file_type in ["txt", "md"]:
            text = content.decode("utf-8", errors="ignore")
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
        
        # Clean and normalize text
        text = self._clean_text(text)
        
        # Split into chunks
        chunks = self._create_chunks(
            text=text,
            resource_id=resource_id,
            workspace_id=workspace_id,
            metadata={"filename": filename, "file_type": file_type}
        )
        
        logger.info(f"Created {len(chunks)} chunks from {filename}")
        
        return chunks
    
    def _extract_pdf(self, content: bytes) -> str:
        """Extract text from PDF"""
        try:
            pdf_file = io.BytesIO(content)
            reader = PdfReader(pdf_file)
            
            text_parts = []
            for page_num, page in enumerate(reader.pages, 1):
                text = page.extract_text()
                if text.strip():
                    # Add page marker for better context
                    text_parts.append(f"[Page {page_num}]\n{text}")
            
            return "\n\n".join(text_parts)
        except Exception as e:
            logger.error(f"PDF extraction failed: {e}")
            raise
    
    def _extract_docx(self, content: bytes) -> str:
        """Extract text from DOCX"""
        try:
            docx_file = io.BytesIO(content)
            document = Document(docx_file)
            
            text_parts = []
            for para in document.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Also extract text from tables
            for table in document.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)
            
            return "\n\n".join(text_parts)
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            raise
    
    def _extract_xlsx(self, content: bytes) -> str:
        """Extract text from XLSX"""
        try:
            xlsx_file = io.BytesIO(content)
            workbook = openpyxl.load_workbook(xlsx_file, data_only=True)
            
            text_parts = []
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                text_parts.append(f"[Sheet: {sheet_name}]")
                
                for row in sheet.iter_rows(values_only=True):
                    row_text = " | ".join(str(cell) for cell in row if cell is not None)
                    if row_text.strip():
                        text_parts.append(row_text)
            
            return "\n".join(text_parts)
        except Exception as e:
            logger.error(f"XLSX extraction failed: {e}")
            raise
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        # Remove excessive whitespace
        text = " ".join(text.split())
        
        # Replace multiple newlines with double newline
        import re
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        # Remove special characters that might interfere with processing
        # (but preserve important punctuation)
        text = re.sub(r'[\x00-\x08\x0b-\x0c\x0e-\x1f\x7f-\x9f]', '', text)
        
        return text.strip()
    
    def _create_chunks(
        self,
        text: str,
        resource_id: UUID,
        workspace_id: UUID,
        metadata: Dict[str, Any]
    ) -> List[Dict[str, Any]]:
        """
        Create intelligent chunks with Parent Document Retrieval support.
        
        Strategy:
        1. Split into large "Parent Docs" (DEFAULT_CHUNK_SIZE, e.g. 1024 tokens)
        2. Split Parents into smaller "Child Chunks" (CHILD_CHUNK_SIZE, e.g. 256 tokens)
        3. Index Children, but store link to Parent.
        """
        chunks = []
        chunk_index = 0
        
        # 1. Create Parent Chunks
        parent_texts = self._split_text(
            text, 
            chunk_size=settings.DEFAULT_CHUNK_SIZE, 
            overlap=settings.CHUNK_OVERLAP
        )
        
        for parent_text in parent_texts:
            parent_chunk_id = uuid4()
            
            # 2. Create Child Chunks from this Parent
            # Use smaller overlap for children to keep them distinct
            child_texts = self._split_text(
                parent_text, 
                chunk_size=settings.CHILD_CHUNK_SIZE, 
                overlap=50 
            )
            
            # If parent is small enough, it might result in just 1 child identical to parent
            
            for child_text in child_texts:
                chunks.append(self._create_chunk_dict(
                    content=child_text,
                    chunk_index=chunk_index,
                    resource_id=resource_id,
                    workspace_id=workspace_id,
                    metadata=metadata,
                    parent_content=parent_text,
                    parent_chunk_id=parent_chunk_id
                ))
                chunk_index += 1
        
        return chunks

    def _split_text(self, text: str, chunk_size: int, overlap: int) -> List[str]:
        """
        Helper to split text into chunks of roughly `chunk_size` tokens.
        Respects paragraphs and sentences.
        """
        chunks = []
        paragraphs = text.split("\n\n")
        
        current_chunk = []
        current_tokens = 0
        
        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
            
            para_tokens = len(self.tokenizer.encode(para))
            
            # If paragraph is massive, split strictly by sentences
            if para_tokens > chunk_size:
                # Flush current chunk first
                if current_chunk:
                    chunks.append(" ".join(current_chunk))
                    # Handle overlap for next chunk
                    overlap_text = self._get_overlap_text(current_chunk, overlap)
                    current_chunk = [overlap_text] if overlap_text else []
                    current_tokens = len(self.tokenizer.encode(overlap_text)) if overlap_text else 0

                # Split complex paragraph
                sentences = self._split_into_sentences(para)
                for sentence in sentences:
                    s_tokens = len(self.tokenizer.encode(sentence))
                    
                    if current_tokens + s_tokens > chunk_size:
                        if current_chunk:
                            chunks.append(" ".join(current_chunk))
                            
                            # Overlap logic
                            overlap_text = self._get_overlap_text(current_chunk, overlap)
                            current_chunk = []
                            if overlap_text:
                                current_chunk.append(overlap_text)
                                current_tokens = len(self.tokenizer.encode(overlap_text))
                            else:
                                current_tokens = 0
                    
                    current_chunk.append(sentence)
                    current_tokens += s_tokens
            
            # Normal paragraph fits or helps fill chunk
            elif current_tokens + para_tokens > chunk_size:
                # Flush
                chunks.append(" ".join(current_chunk))
                
                # Start new with overlap
                overlap_text = self._get_overlap_text(current_chunk, overlap)
                current_chunk = []
                if overlap_text:
                    current_chunk.append(overlap_text)
                    current_tokens = len(self.tokenizer.encode(overlap_text))
                else:
                    current_tokens = 0
                
                current_chunk.append(para)
                current_tokens += para_tokens
            else:
                current_chunk.append(para)
                current_tokens += para_tokens
        
        # Add final chunk
        if current_chunk:
            chunks.append(" ".join(current_chunk))
            
        return chunks

    def _get_overlap_text(self, current_chunk: List[str], overlap_tokens: int) -> str:
        """Calculate overlap text from the end of current_chunk"""
        if not current_chunk:
            return ""
        
        overlap_text = ""
        current_overlap_tokens = 0
        
        # Traverse backwards
        for part in reversed(current_chunk):
            part_tokens = len(self.tokenizer.encode(part))
            if current_overlap_tokens + part_tokens <= overlap_tokens:
                overlap_text = part + " " + overlap_text
                current_overlap_tokens += part_tokens
            else:
                break
        return overlap_text.strip()
    
    def _split_into_sentences(self, text: str) -> List[str]:
        """Split text into sentences"""
        import re
        
        # Simple sentence splitter
        # In production, consider using spaCy or nltk for better accuracy
        sentences = re.split(r'(?<=[.!?])\s+', text)
        return [s.strip() for s in sentences if s.strip()]
    
    def _create_chunk_dict(
        self,
        content: str,
        chunk_index: int,
        resource_id: UUID,
        workspace_id: UUID,
        metadata: Dict[str, Any],
        parent_content: str = None,
        parent_chunk_id: UUID = None
    ) -> Dict[str, Any]:
        """Create a chunk dictionary"""
        token_count = len(self.tokenizer.encode(content))
        
        return {
            "id": uuid4(),
            "content": content,
            "parent_content": parent_content,
            "parent_chunk_id": parent_chunk_id,
            "chunk_index": chunk_index,
            "token_count": token_count,
            "resource_id": resource_id,
            "workspace_id": workspace_id,
            "metadata": {
                **metadata,
                "chunk_index": chunk_index,
                "token_count": token_count
            }
        }
     
class DoclingIngestionService(IngestionService):
    """Customized Ingestion Service for Docling-specific needs"""
    
    def __init__(self):
        super().__init__()
        pipeline_options = PdfPipelineOptions()
        pipeline_options.do_ocr = True 
        pipeline_options.do_table_structure = False
        self.converter = DocumentConverter(
            allowed_formats=[
                InputFormat.PDF,
                InputFormat.DOCX,
                InputFormat.PPTX,
                InputFormat.IMAGE,
                InputFormat.HTML,
                InputFormat.MD,
                InputFormat.XLSX,
                InputFormat.CSV,
            ],
            format_options={
            InputFormat.PDF: PdfFormatOption(pipeline_cls=StandardPdfPipeline, backend=PyPdfiumDocumentBackend),
            InputFormat.DOCX: WordFormatOption(pipeline_cls=SimplePipeline, backend=MsWordDocumentBackend),
            InputFormat.CSV: CsvFormatOption(pipeline_cls=SimplePipeline, backend=CsvDocumentBackend),
            InputFormat.XLSX: ExcelFormatOption(pipeline_cls=SimplePipeline, backend=MsExcelDocumentBackend),
            InputFormat.PPTX: PowerpointFormatOption(pipeline_cls=SimplePipeline, backend=MsPowerpointDocumentBackend),
            InputFormat.MD: MarkdownFormatOption(pipeline_cls=SimplePipeline, backend=MarkdownDocumentBackend),
            InputFormat.HTML: HTMLFormatOption(pipeline_cls=SimplePipeline, backend=HTMLDocumentBackend),
            InputFormat.IMAGE: ImageFormatOption(pipeline_cls=SimplePipeline, backend=ImageDocumentBackend),
            },
        )
        
    async def process_document(
        self,
        content: bytes,
        filename: str,
        file_type: str,
        resource_id: UUID,
        workspace_id: UUID
    ) -> List[Dict[str, Any]]:
        """
        Process a document into chunks using Docling's DocumentConverter
        """
        logger.info(f"Processing document with Docling: {filename} ({file_type})")
        
        supported_types = [
            "pdf", "docx", "doc", "pptx", "ppt",
            "xlsx", "xls", "txt", "md", "html", "htm",
            "png", "jpg", "jpeg", "tiff", "bmp"
        ]
        
        if file_type.lower() not in supported_types:
            raise ValueError(f"Unsupported file type: {file_type}")
        
        try:
            text = self._extract_with_docling(content, filename, file_type)

            # Clean and normalize text
            text = self._clean_text(text)

            # Split into chunks (inherits logic from IngestionService)
            chunks = self._create_chunks(
                text=text,
                resource_id=resource_id,
                workspace_id=workspace_id,
                metadata={"filename": filename, "file_type": file_type}
            )

            logger.info(f"Created {len(chunks)} chunks from {filename} using Docling")

            return chunks

        except Exception as e:
            logger.error(f"Docling failed for {filename}, falling back: {e}")

            # Fallback using super's process_document
            return await super().process_document(
                content=content,
                filename=filename,
                file_type=file_type,
                resource_id=resource_id,
                workspace_id=workspace_id
            )

    def _extract_with_docling(self, content: bytes, filename: str, file_type: str) -> str:
        """
        Extract text using Docling for superior quality and format support
        """
        # Create temporary file (Docling works with file paths)
        with tempfile.NamedTemporaryFile(delete=False, suffix=f".{file_type}") as tmp_file:
            tmp_file.write(content)
            tmp_path = tmp_file.name
        
        try:
            # Convert document
            conv_res = self.converter.convert(tmp_path)
            
            # Extract text from the document
            text_parts = []
            
            # Get document structure
            doc = conv_res.document
            
            # Extract main text content
            if hasattr(doc, 'export_to_markdown'):
                # Export to markdown preserves structure better
                text_parts.append(doc.export_to_markdown())
            else:
                # Fallback: extract text from document body
                for element in doc.body:
                    if hasattr(element, 'text') and element.text:
                        text_parts.append(element.text)
            
            # Extract tables separately for better formatting
            if hasattr(doc, 'tables') and doc.tables:
                for i, table in enumerate(doc.tables):
                    text_parts.append(f"\n[Table {i+1}]")
                    if hasattr(table, 'export_to_markdown'):
                        text_parts.append(table.export_to_markdown())
                    elif hasattr(table, 'data'):
                        # Format table data as text
                        for row in table.data:
                            text_parts.append(" | ".join(str(cell) for cell in row))
            
            return "\n\n".join(text_parts)
            
        finally:
            # Clean up temporary file
            try:
                os.unlink(tmp_path)
            except Exception as e:
                logger.warning(f"Failed to delete temporary file {tmp_path}: {e}")
   